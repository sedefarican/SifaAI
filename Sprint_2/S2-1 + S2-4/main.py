from prompt_engine import generate_tahlil_prompt, generate_reminder_prompt
from utils import send_prompt_to_gemini
from docx import Document
from datetime import datetime

# Word dökümanı oluştur
document = Document()
document.add_heading("ŞifaAI - Yapay Zeka Yanıtları", 0)
document.add_paragraph(f"Tarih: {datetime.today().strftime('%d.%m.%Y')}")

# Örnek kullanıcı bilgisi
kullanici = {
    "isim": "Ayşe Yılmaz",
    "yaş": 65,
    "cinsiyet": "Kadın",
    "kilo": 74,
    "boy": 158,
    "kronikHastalıklar": ["Tip 2 Diyabet", "Hipertansiyon"],
    "ilaçlar": ["Metformin", "Amlodipin"],
    "alerjiler": ["Penisilin"]
}

# Örnek tahlil verisi
tahlil_metni = """
- WBC: 8,11 (Referans: 4.8 - 10.8)
- LY%: 39,7 (Referans: 5.5 - 20.5)
- EO%: 3,9 (Referans: 0.0 - 2.0)
- HCT: 50,4 (Referans: 38.0 - 42.0)
- MCH: 23,6 (Referans: 27.0 - 31.0)
"""

# --- Tahlil Yorumlama ---
print("Tahlil Yorumu:\n")
tahlil_promptu = generate_tahlil_prompt(kullanici, tahlil_metni)
tahlil_cevap = send_prompt_to_gemini(tahlil_promptu)
print(tahlil_cevap)

#tahlil yorumunu word dosyasina kaydetme
'''document.add_heading("1. Tahlil Yorumu", level=1)
document.add_paragraph(tahlil_cevap)

# --- İlaç Hatırlatma ---
print("\nİlaç Hatırlatma:\n")
hatirlatma_promptu = generate_reminder_prompt(kullanici)
hatirlatma_cevap = send_prompt_to_gemini(hatirlatma_promptu)
print(hatirlatma_cevap)

document.add_heading("2. İlaç ve Randevu Hatırlatma", level=1)
document.add_paragraph(hatirlatma_cevap)

# --- Dosyayı kaydet ---
output_path = "AI_Yanitlari.docx"
document.save(output_path)
print(f"\nYanıtlar Word dosyasına kaydedildi: {output_path}")'''
